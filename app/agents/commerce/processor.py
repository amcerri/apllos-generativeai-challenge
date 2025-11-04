"""
Commerce document processor (multi-format text extraction with OCR).

Overview
--------
Extract text from various document formats (PDF, DOCX, TXT, images) with
OCR fallback for scanned documents. This processor handles binary attachments
and converts them to clean text suitable for LLM processing. It supports
multiple extraction methods with graceful fallbacks.

Design
------
- Input: binary content, filename, and MIME type from attachments
- Output: extracted text, metadata, and processing method used
- Multi-format support: PDF (pypdf + OCR), DOCX (python-docx), TXT (direct), images (OCR)
- OCR fallback: when text extraction fails or produces poor results
- Dependency-light: optional imports with graceful degradation

Integration
-----------
- Used by the commerce agent before the LLM extractor
- Logging uses centralized infra if available; tracing spans are optional
- Returns structured data compatible with existing commerce pipeline

Usage
-----
>>> from app.agents.commerce.processor import DocumentProcessor
>>> processor = DocumentProcessor()
>>> result = processor.process_attachment({
...     "filename": "invoice.pdf",
...     "content": b"...",  # binary content
...     "mime_type": "application/pdf"
... })
>>> result["text"], result["method"]
('Invoice #INV-001...', 'pdf_direct')
"""

from __future__ import annotations

import io
import os
import re
import tempfile
from pathlib import Path
from typing import Any, Final

try:  # Optional logger
    from app.infra.logging import get_logger
except Exception:  # pragma: no cover - optional
    import logging as _logging

    def get_logger(component: str, **initial_values: Any) -> Any:
        return _logging.getLogger(component)


# Tracing (optional; single alias)
start_span: Any
try:
    from app.infra.tracing import start_span as _start_span
    start_span = _start_span
except Exception:  # pragma: no cover - optional
    from contextlib import nullcontext as _nullcontext

    def _fallback_start_span(_name: str, _attrs: dict[str, Any] | None = None):
        return _nullcontext()

    start_span = _fallback_start_span


# Optional dependencies with graceful fallbacks
PdfReader: Any = None
try:
    from pypdf import PdfReader as _PdfReader
    PdfReader = _PdfReader
except Exception:  # pragma: no cover - optional
    pass

Document: Any = None
try:
    from docx import Document as _Document
    Document = _Document
except Exception:  # pragma: no cover - optional
    pass

pytesseract: Any = None
try:
    import pytesseract as _pytesseract
    pytesseract = _pytesseract
except Exception:  # pragma: no cover - optional
    pass

Image: Any = None
try:
    from PIL import Image as _Image
    Image = _Image
except Exception:  # pragma: no cover - optional
    pass

pdf2image: Any = None
try:
    import pdf2image as _pdf2image
    pdf2image = _pdf2image
except Exception:  # pragma: no cover - optional
    pass

# Optional OpenCV for advanced preprocessing
cv2: Any = None
try:  # pragma: no cover - optional
    import cv2 as _cv2  # type: ignore
    cv2 = _cv2
except Exception:
    cv2 = None


__all__ = ["DocumentProcessor"]


# ---------------------------------------------------------------------------
# Document Processor
# ---------------------------------------------------------------------------
class DocumentProcessor:
    """Extract text from various document formats with OCR fallback."""

    # Minimum text length to consider extraction successful
    MIN_TEXT_LENGTH: Final[int] = 50
    
    # OCR confidence threshold (0-100)
    OCR_CONFIDENCE_THRESHOLD: Final[int] = 60

    def __init__(self) -> None:
        self.log = get_logger("agent.commerce.processor")

    def process_attachment(self, attachment: dict[str, Any]) -> dict[str, Any]:
        """Process attachment and extract text.

        Parameters
        ----------
        attachment: Dict with 'filename', 'content' (bytes or base64 string), and optional 'mime_type'

        Returns
        -------
        Dict with 'text', 'metadata', 'method', 'warnings', and 'success'
        """
        filename = attachment.get("filename", "unknown")
        self.log.info("Starting document processing", extra={"file_name": filename, "mime_type": attachment.get("mime_type")})
        
        with start_span("agent.commerce.process_attachment"):
            filename = attachment.get("filename", "unknown")
            content = attachment.get("content", b"")
            mime_type = attachment.get("mime_type", "")
            
            if isinstance(content, str):
                # Check if it's base64 encoded binary content
                try:
                    import base64
                    # Try to decode as base64
                    decoded_content = base64.b64decode(content)
                    # If successful and looks like binary (not just ASCII), use decoded content
                    if len(decoded_content) > 0 and not all(32 <= b <= 126 for b in decoded_content[:100]):
                        content = decoded_content
                    else:
                        # It's actually text content, not base64
                        return {
                            "text": content,
                            "metadata": {"filename": filename, "size": len(content)},
                            "method": "text_direct",
                            "warnings": [],
                            "success": True
                        }
                except Exception:
                    # Not base64, treat as text
                    return {
                        "text": content,
                        "metadata": {"filename": filename, "size": len(content)},
                        "method": "text_direct",
                        "warnings": [],
                        "success": True
                    }
            
            if not isinstance(content, (bytes, bytearray)):
                return self._error_result("Invalid content type", filename)
            
            # Determine file type
            file_ext = Path(filename).suffix.lower()
            
            # Try appropriate extraction method
            if file_ext == ".pdf" or "pdf" in mime_type.lower():
                return self._process_pdf(content, filename)
            elif file_ext == ".docx" or "wordprocessingml" in mime_type.lower():
                return self._process_docx(content, filename)
            elif file_ext in {".txt", ".md"} or "text" in mime_type.lower():
                return self._process_text(content, filename)
            elif file_ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"} or "image" in mime_type.lower():
                return self._process_image(content, filename)
            else:
                # Unknown format, try OCR as last resort
                self.log.warning("Unknown file format, attempting OCR", extra={"file_name": filename, "mime_type": mime_type})
                return self._process_image(content, filename)

    def _is_text_useful(self, text: str) -> bool:
        """Check if extracted text is useful content or just metadata.
        
        Returns False if text appears to be only metadata (URLs, dates, filenames).
        """
        if not text or len(text.strip()) < self.MIN_TEXT_LENGTH:
            return False
        
        text_lower = text.lower()
        words = text.split()
        
        # Heuristics to detect metadata-only content
        # 1. Too many URLs
        url_count = text_lower.count('http://') + text_lower.count('https://') + text_lower.count('www.')
        if url_count > 0 and len(words) < 20:
            return False
        
        # 2. Mostly dates/timestamps (format: DD/MM/YYYY, MM/DD/YYYY, etc.)
        date_pattern = r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}'
        dates = re.findall(date_pattern, text)
        if len(dates) > len(words) * 0.3:  # More than 30% of content is dates
            return False
        
        # 3. Too many file paths/extensions
        path_pattern = r'[\w\-_]+\.(pdf|png|jpg|jpeg|docx|doc|txt|zip)'
        paths = re.findall(path_pattern, text_lower)
        if len(paths) > len(words) * 0.2:  # More than 20% of content is file references
            return False
        
        # 4. Mostly dimensions (e.g., "720×1040", "1280x720")
        dimension_pattern = r'\d+\s*[×x]\s*\d+'
        dimensions = re.findall(dimension_pattern, text_lower)
        if len(dimensions) > 0 and len(words) < 15:
            return False
        
        # 5. Very short average word length (likely metadata fragments)
        if words:
            avg_word_len = sum(len(w) for w in words) / len(words)
            if avg_word_len < 3.0 and len(words) < 30:
                return False
        
        return True

    def _process_pdf(self, content: bytes, filename: str) -> dict[str, Any]:
        """Extract text from PDF with OCR fallback."""
        warnings = []
        
        # Try direct PDF text extraction first
        if PdfReader is not None:
            try:
                with io.BytesIO(content) as pdf_buffer:
                    reader = PdfReader(pdf_buffer)
                    texts = []
                    
                    for i, page in enumerate(reader.pages):
                        try:
                            page_text = page.extract_text() or ""
                            texts.append(page_text)
                        except Exception as e:
                            warnings.append(f"Failed to extract page {i+1}: {str(e)}")
                            continue
                    
                    full_text = "\n".join(texts).strip()
                    
                    # Check if text is useful (not just metadata)
                    if len(full_text) >= self.MIN_TEXT_LENGTH and self._is_text_useful(full_text):
                        return {
                            "text": full_text,
                            "metadata": {
                                "filename": filename,
                                "pages": len(reader.pages),
                                "size": len(content)
                            },
                            "method": "pdf_direct",
                            "warnings": warnings,
                            "success": True
                        }
                    else:
                        if len(full_text) >= self.MIN_TEXT_LENGTH:
                            warnings.append("PDF text extraction yielded metadata-only content, trying OCR")
                        else:
                            warnings.append("PDF text extraction yielded insufficient text, trying OCR")
                        
            except Exception as e:
                warnings.append(f"PDF direct extraction failed: {str(e)}")
        else:
            warnings.append("pypdf not available, trying OCR")
        
        # Fallback to OCR
        return self._process_pdf_ocr(content, filename, warnings)

    def _process_pdf_ocr(self, content: bytes, filename: str, warnings: list[str]) -> dict[str, Any]:
        """Extract text from PDF using OCR."""
        if pdf2image is None or pytesseract is None or Image is None:
            warnings.append("OCR dependencies not available")
            return self._error_result("Cannot process PDF: no OCR support", filename, warnings)
        
        try:
            self.log.info("Starting OCR processing for PDF", extra={"file_name": filename, "size_bytes": len(content)})
            
            # Convert PDF to images
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
                tmp_pdf.write(content)
                tmp_pdf.flush()
                
                try:
                    self.log.info("Converting PDF to images...", extra={"file_name": filename})
                    images = pdf2image.convert_from_path(tmp_pdf.name, dpi=200)
                    self.log.info(f"PDF converted to {len(images)} image(s), starting OCR...", extra={"file_name": filename, "pages": len(images)})
                    
                    texts = []
                    
                    for i, image in enumerate(images):
                        try:
                            self.log.debug(f"Processing page {i+1}/{len(images)} with OCR...", extra={"file_name": filename, "page": i+1})
                            # OCR each page
                            page_text = pytesseract.image_to_string(image, lang='eng+por')
                            if page_text.strip():
                                texts.append(page_text.strip())
                                self.log.debug(f"Page {i+1} OCR completed, extracted {len(page_text)} characters", extra={"file_name": filename, "page": i+1})
                        except Exception as e:
                            warnings.append(f"OCR failed for page {i+1}: {str(e)}")
                            self.log.warning(f"OCR failed for page {i+1}", extra={"file_name": filename, "page": i+1, "error": str(e)})
                            continue
                    
                    full_text = "\n".join(texts).strip()
                    self.log.info("OCR processing completed", extra={"file_name": filename, "total_text_length": len(full_text), "pages_processed": len(images)})
                    
                    return {
                        "text": full_text,
                        "metadata": {
                            "filename": filename,
                            "pages": len(images),
                            "size": len(content)
                        },
                        "method": "pdf_ocr",
                        "warnings": warnings,
                        "success": len(full_text) >= self.MIN_TEXT_LENGTH
                    }
                    
                finally:
                    # Clean up temp file
                    try:
                        os.unlink(tmp_pdf.name)
                    except Exception:
                        pass
                        
        except Exception as e:
            error_msg = str(e)
            # Check for common OCR setup issues
            if "poppler" in error_msg.lower() or "pdfinfo" in error_msg.lower() or "PDFInfoNotInstalledError" in error_msg:
                detailed_error = (
                    "PDF OCR requires Poppler utilities to be installed. "
                    "On macOS: 'brew install poppler'. "
                    "On Ubuntu/Debian: 'apt-get install poppler-utils'. "
                    "On Windows: Download from poppler.freedesktop.org"
                )
                warnings.append(f"PDF OCR setup issue: {detailed_error}")
                self.log.warning("OCR not available - Poppler missing", extra={"file_name": filename, "error": error_msg})
            elif "tesseract" in error_msg.lower():
                detailed_error = (
                    "Tesseract OCR engine not found. "
                    "On macOS: 'brew install tesseract tesseract-lang'. "
                    "On Ubuntu/Debian: 'apt-get install tesseract-ocr tesseract-ocr-eng tesseract-ocr-por'"
                )
                warnings.append(f"OCR setup issue: {detailed_error}")
                self.log.warning("OCR not available - Tesseract missing", extra={"file_name": filename, "error": error_msg})
            else:
                warnings.append(f"PDF OCR failed: {error_msg}")
            
            # Return the metadata text as fallback if we have it, even if it's not useful
            # This allows the LLM to at least see something and potentially extract structure
            try:
                if PdfReader is not None:
                    with io.BytesIO(content) as pdf_buffer:
                        reader = PdfReader(pdf_buffer)
                        fallback_texts = []
                        for page in reader.pages:
                            try:
                                page_text = page.extract_text() or ""
                                if page_text.strip():
                                    fallback_texts.append(page_text.strip())
                            except Exception:
                                continue
                        fallback_text = "\n".join(fallback_texts).strip()
                        if fallback_text:
                            warnings.append("Using fallback text extraction (may be metadata only)")
                            return {
                                "text": fallback_text,
                                "metadata": {
                                    "filename": filename,
                                    "pages": len(reader.pages),
                                    "size": len(content),
                                    "ocr_failed": True
                                },
                                "method": "pdf_direct_fallback",
                                "warnings": warnings,
                                "success": len(fallback_text) >= 10  # Lower threshold for fallback
                            }
            except Exception:
                pass
            
            return self._error_result("PDF OCR processing failed", filename, warnings)

    def _process_docx(self, content: bytes, filename: str) -> dict[str, Any]:
        """Extract text from DOCX file."""
        if Document is None:
            return self._error_result("python-docx not available", filename)
        
        try:
            with io.BytesIO(content) as docx_buffer:
                doc = Document(docx_buffer)
                
                # Extract text from paragraphs
                paragraphs = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(text)
                
                # Extract text from tables
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            tables_text.append(" | ".join(row_text))
                
                # Combine all text
                all_text = []
                if paragraphs:
                    all_text.extend(paragraphs)
                if tables_text:
                    all_text.append("\n--- Tables ---")
                    all_text.extend(tables_text)
                
                full_text = "\n".join(all_text).strip()
                
                return {
                    "text": full_text,
                    "metadata": {
                        "filename": filename,
                        "paragraphs": len(paragraphs),
                        "tables": len(doc.tables),
                        "size": len(content)
                    },
                    "method": "docx_direct",
                    "warnings": [],
                    "success": len(full_text) >= self.MIN_TEXT_LENGTH
                }
                
        except Exception as e:
            return self._error_result(f"DOCX processing failed: {str(e)}", filename)

    def _process_text(self, content: bytes, filename: str) -> dict[str, Any]:
        """Extract text from plain text file."""
        try:
            # Try UTF-8 first, then fallback encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    return {
                        "text": text,
                        "metadata": {
                            "filename": filename,
                            "encoding": encoding,
                            "size": len(content)
                        },
                        "method": "text_direct",
                        "warnings": [],
                        "success": True
                    }
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use errors='ignore'
            text = content.decode('utf-8', errors='ignore')
            return {
                "text": text,
                "metadata": {
                    "filename": filename,
                    "encoding": "utf-8-ignore",
                    "size": len(content)
                },
                "method": "text_fallback",
                "warnings": ["Used fallback encoding, some characters may be lost"],
                "success": True
            }
            
        except Exception as e:
            return self._error_result(f"Text processing failed: {str(e)}", filename)

    def _process_image(self, content: bytes, filename: str) -> dict[str, Any]:
        """Extract text from image using OCR."""
        if pytesseract is None or Image is None:
            return self._error_result("OCR dependencies not available", filename)
        
        try:
            # Load image
            with io.BytesIO(content) as img_buffer:
                image = Image.open(img_buffer)
                
                # Convert to RGB if necessary
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                # Advanced preprocessing (optional OpenCV pipeline)
                pre_meta: dict[str, Any] = {}
                if cv2 is not None:
                    try:
                        import numpy as _np  # type: ignore
                        # PIL -> OpenCV BGR
                        cv_img = cv2.cvtColor(_np.array(image), cv2.COLOR_RGB2BGR)
                        gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
                        # Binarize (Otsu)
                        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                        # Deskew (estimate angle via moments)
                        coords = _np.column_stack(_np.where(thresh > 0))
                        angle = 0.0
                        if coords.size > 0:
                            rect = cv2.minAreaRect(coords)
                            angle = rect[-1]
                            if angle < -45:
                                angle = -(90 + angle)
                            else:
                                angle = -angle
                            (h, w) = thresh.shape[:2]
                            M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
                            thresh = cv2.warpAffine(thresh, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
                        pre_meta.update({"deskew_angle": float(angle)})
                        # Back to PIL for pytesseract
                        image = Image.fromarray(cv2.cvtColor(thresh, cv2.COLOR_GRAY2RGB))
                    except Exception:
                        # Fallback to original image on failure
                        pass

                # OCR with both English and Portuguese (configurable via env)
                lang = os.getenv("OCR_LANG", "eng+por")
                text = pytesseract.image_to_string(image, lang=lang)
                
                # Get confidence data if available
                confidence = None
                try:
                    data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT, lang=lang)
                    confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                    if confidences:
                        confidence = sum(confidences) / len(confidences)
                except Exception:
                    pass
                
                warnings = []
                if confidence is not None and confidence < self.OCR_CONFIDENCE_THRESHOLD:
                    warnings.append(f"Low OCR confidence: {confidence:.1f}%")
                
                return {
                    "text": text.strip(),
                    "metadata": {
                        "filename": filename,
                        "image_size": image.size,
                        "image_mode": image.mode,
                        "ocr_confidence": confidence,
                        "preprocess": pre_meta,
                        "size": len(content)
                    },
                    "method": "image_ocr",
                    "warnings": warnings,
                    "success": len(text.strip()) >= self.MIN_TEXT_LENGTH
                }
                
        except Exception as e:
            return self._error_result(f"Image OCR failed: {str(e)}", filename)

    def _error_result(self, error: str, filename: str, warnings: list[str] | None = None) -> dict[str, Any]:
        """Return standardized error result."""
        self.log.error("Document processing failed", extra={"file_name": filename, "error": error})
        return {
            "text": "",
            "metadata": {"filename": filename, "error": error},
            "method": "failed",
            "warnings": (warnings or []) + [error],
            "success": False
        }
